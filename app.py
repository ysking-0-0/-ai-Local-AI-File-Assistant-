import os, json, uuid, time, argparse, shutil, threading
from pathlib import Path
from datetime import datetime
from typing import List, Optional

import chromadb
from chromadb import Documents, EmbeddingFunction, Embeddings
from fastapi import FastAPI, UploadFile, File, Query, HTTPException, Form
from fastapi.responses import JSONResponse, FileResponse, Response
import uvicorn
import gradio as gr
from openai import OpenAI
from sentence_transformers import SentenceTransformer
import docx  # Word 文档解析

# ==================== 配置区 ====================
DEEPSEEK_API_KEY = "sk-your-deepseek-api-key-here"          # ← 替换成你的新 Key
PROJECT_DIR = Path(__file__).parent.resolve()
CHROMA_DIR = PROJECT_DIR / "chroma_db"
STORAGE_DIR = PROJECT_DIR / "我的文件"
TEMP_DIR = PROJECT_DIR / "temp_uploads"
FILE_THRESHOLD = 10 * 1024 * 1024          # 10 MB

CATEGORY_RULES = {
    "0_文本笔记": {"types": ["文本笔记"], "tags": ["文档", "教程", "项目说明"]},
    "1_配置密钥": {"types": ["配置密钥"], "tags": ["API Key", "配置"]},
    "2_项目文档": {"types": ["项目文档"], "tags": ["项目说明", "文档", "教程"]},
    "3_项目截图": {"types": ["项目截图"], "tags": ["运行截图"]},
    "大文件存储": {"types": [], "tags": []}
}

FILE_TYPES = ["文本笔记", "配置密钥", "项目文档", "项目截图"]
FILE_TAGS = ["API Key", "配置", "文档", "教程", "项目说明", "运行截图"]

# 确保目录存在
STORAGE_DIR.mkdir(parents=True, exist_ok=True)
TEMP_DIR.mkdir(parents=True, exist_ok=True)

# ==================== 全局变量 ====================
NGROK_PUBLIC_URL = None  # ngrok 公网地址，启动时赋值

# 下载链接的基础地址（Gradio 运行时动态设置）
_DOWNLOAD_BASE = "http://127.0.0.1:8000"

# ==================== 嵌入模型 ====================
class LocalEmbeddingFunction(EmbeddingFunction):
    def __init__(self, model_name: str = "paraphrase-multilingual-MiniLM-L12-v2"):
        print(f"正在加载本地嵌入模型: {model_name}")
        # 先尝试从本地缓存加载（模型已预先下载到 HuggingFace 缓存）
        try:
            self.model = SentenceTransformer(model_name, local_files_only=True)
        except Exception:
            print("[WARN] 本地缓存未找到，尝试在线下载...")
            self.model = SentenceTransformer(model_name)
        print("模型加载完成。")

    def __call__(self, input: Documents) -> Embeddings:
        embeddings = self.model.encode(list(input), convert_to_numpy=True)
        return embeddings.tolist()

# ==================== 初始化客户端 ====================
openai_client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com/v1")
embedding_func = LocalEmbeddingFunction()
chroma_client = chromadb.PersistentClient(path=str(CHROMA_DIR))
collection = chroma_client.get_or_create_collection(
    name="my_kb",
    embedding_function=embedding_func
)

# ==================== 核心函数 ====================

def extract_text_content(file_path: Path) -> str:
    """提取文件文本内容"""
    file_type = file_path.suffix.lower()
    file_size = file_path.stat().st_size

    if file_size > FILE_THRESHOLD:
        return f"大文件(>10MB)，路径: {file_path}"

    try:
        if file_type in ['.txt', '.md', '.py', '.json', '.yaml', '.yml', '.ini', '.cfg', '.log']:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif file_type == '.docx':
            doc = docx.Document(str(file_path))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            content = '\n'.join(paragraphs)
            if not content:
                tables_text = []
                for table in doc.tables:
                    for row in table.rows:
                        cells = [cell.text for cell in row.cells]
                        tables_text.append(' | '.join(cells))
                content = '\n'.join(tables_text)
            return content if content else f"Word文档无文字内容，路径: {file_path}"
        elif file_type == '.pdf':
            return f"PDF文件，路径: {file_path}"
        else:
            return f"不支持直接读取格式，路径: {file_path}"
    except Exception as e:
        return f"文件读取失败，路径: {file_path}，错误: {str(e)}"


def get_file_sample(file_path: Path) -> str:
    """提取文件前500字符用于分类判断"""
    try:
        # 先尝试直接读取文本内容的前500字符
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read(500)
    except Exception:
        if file_path.suffix.lower() == '.docx':
            try:
                doc = docx.Document(str(file_path))
                sample = '\n'.join([p.text for p in doc.paragraphs[:3] if p.text.strip()])
                return sample if sample else f"[Word文档: {file_path.name}]"
            except:
                return f"[非文本文件: {file_path.name}]"
        else:
            return f"[非文本文件: {file_path.name}]"


def classify_file(file_path: Path) -> dict:
    """使用 DeepSeek 自动分类文件，返回 {type, tags, summary}"""
    sample = get_file_sample(file_path)

    prompt = f"""你是一个文件分类专家。根据以下规则分类文件：
{json.dumps(CATEGORY_RULES, ensure_ascii=False)}
文件路径：{file_path}
内容片段：{sample}
请输出 JSON：{{"type":"类型","tags":["标签"],"summary":"一句话摘要"}}"""

    resp = openai_client.chat.completions.create(
        model="deepseek-chat",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1
    )
    result = resp.choices[0].message.content.strip().strip("```json").strip("```")
    return json.loads(result)


def index_file(file_path: Path, classification: dict) -> str:
    """将文件存入向量库，返回 chroma_id"""
    content = extract_text_content(file_path)
    file_size = file_path.stat().st_size
    filename = file_path.name
    now_iso = datetime.now().isoformat()

    doc_id = str(uuid.uuid4())
    collection.add(
        ids=[doc_id],
        documents=[content],
        metadatas=[{
            "type": classification.get('type', '未知'),
            "tags": ", ".join(classification.get('tags', [])),
            "summary": classification.get('summary', ''),
            "file_path": os.path.normpath(str(file_path)),
            "filename": filename,
            "upload_time": now_iso,
            "file_size": file_size,
            "is_doc_note": "false"   # 标记是否为文本转文档生成
        }]
    )
    return doc_id


def update_metadata(file_id: str, file_type: str, tags: str, summary: str) -> bool:
    """更新 ChromaDB 中文档的元数据"""
    try:
        collection.update(
            ids=[file_id],
            metadatas=[{
                "type": file_type,
                "tags": tags,
                "summary": summary
            }]
        )
        return True
    except Exception as e:
        print(f"更新失败: {e}")
        return False


def delete_document(file_id: str) -> bool:
    """从 ChromaDB 删除文档记录"""
    try:
        collection.delete(ids=[file_id])
        return True
    except Exception as e:
        print(f"删除失败: {e}")
        return False


def get_all_files() -> dict:
    """获取所有已索引文件"""
    all_data = collection.get(include=['metadatas'])
    if not all_data['ids']:
        return {"ids": [], "metadatas": []}
    return all_data


def public_url_for(path: str) -> str:
    """将本地路径转为公网可访问的下载 URL（如果 ngrok 已启动）"""
    if NGROK_PUBLIC_URL is None:
        return None
    # 根据 file_path 查找 chroma_id
    result = collection.get(
        where={"file_path": path},
        include=['metadatas']
    )
    if result['ids']:
        fid = result['ids'][0]
        return f"{NGROK_PUBLIC_URL}/download/{fid}"
    return None


def text_to_doc(text: str, label: str = "") -> dict:
    """将一段文字转为 .md 文档并入库

    返回: {success, doc_id, filename, title, summary}
    """
    try:
        # 1. 用 DeepSeek 生成标题和摘要
        prompt = f"""你是一个文档整理助手。根据以下文字内容，生成一个简短的标题和一句话摘要。
文字内容：
{text[:2000]}

请输出 JSON：{{"title":"标题","summary":"一句话摘要"}}"""

        resp = openai_client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3
        )
        result = resp.choices[0].message.content.strip().strip("```json").strip("```")
        doc_info = json.loads(result)

        title = doc_info.get("title", "未命名文档")
        summary = doc_info.get("summary", "")

        # 2. 生成 .md 文件
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        label_part = f"【{label}】" if label else ""
        safe_title = "".join(c for c in title if c.isalnum() or c in ' -_（）()').strip()
        filename = f"笔记_{label_part}_{safe_title}_{timestamp}.md" if not label else f"{label}_{safe_title}_{timestamp}.md"
        # 清理多余字符
        filename = filename.replace(" ", "_").replace("__", "_")

        md_content = f"""---
title: {title}
label: {label or "未标注"}
created: {datetime.now().isoformat()}
type: 文本笔记
---

# {title}

> **标记**: {label or "无"}
> **生成时间**: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
> **摘要**: {summary}

---

## 原文内容

{text}

---

## 自动摘要

{summary}
"""
        file_path = STORAGE_DIR / filename
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(md_content)

        # 3. 入库
        classification = {
            "type": "文本笔记",
            "tags": ["文档"] + ([label] if label and label not in ["文档"] else []),
            "summary": f"【{label}】{summary}" if label else summary
        }
        doc_id = index_file(file_path, classification)

        return {
            "success": True,
            "doc_id": doc_id,
            "filename": filename,
            "title": title,
            "summary": summary
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


# ==================== 公网暴露 (ngrok) ====================

def start_ngrok(port: int = 8000):
    """启动 ngrok 隧道暴露本地端口"""
    global NGROK_PUBLIC_URL
    try:
        from pyngrok import ngrok

        # 检查是否有 token 配置
        ngrok_token = os.environ.get("NGROK_TOKEN", "")
        if ngrok_token:
            ngrok.set_auth_token(ngrok_token)
            print("已配置 ngrok 认证令牌")

        # 启动隧道
        tunnel = ngrok.connect(port, "http")
        NGROK_PUBLIC_URL = tunnel.public_url
        print(f"\n{'='*60}")
        print(f"  [ngrok] 公网地址: {NGROK_PUBLIC_URL}")
        print(f"  [ngrok] API 端点: {NGROK_PUBLIC_URL}/docs")
        print(f"{'='*60}\n")
        return NGROK_PUBLIC_URL
    except ImportError:
        print("\n[WARN] 未安装 pyngrok，无法自动暴露公网。")
        print("   安装: pip install pyngrok")
        print("   然后: 注册 https://ngrok.com 获取 token")
        print("   设置: set NGROK_TOKEN=你的token")
        print("   启动: python app.py --expose\n")
        return None
    except Exception as e:
        print(f"\n[WARN] ngrok 启动失败: {e}")
        print("   可手动运行: ngrok http 8000\n")
        return None


# ==================== FastAPI ====================

app = FastAPI(title="MyKB - 私人知识库", version="2.0")


@app.get("/")
async def root():
    return {
        "name": "MyKB 私人知识库",
        "version": "2.0",
        "endpoints": {
            "upload": "POST /upload (多文件上传并入库)",
            "upload_preview": "POST /upload-preview (上传预览，不入库)",
            "confirm_upload": "POST /confirm-upload (确认预览结果并入库)",
            "search": "GET /search?q=关键词",
            "download": "GET /download/{file_id}",
            "files": "GET /files (列出所有文件)",
            "update_file": "PUT /files/{file_id} (更新元数据)",
            "delete_file": "DELETE /files/{file_id}",
            "text_to_doc": "POST /text-to-doc (文字转文档)"
        },
        "public_url": NGROK_PUBLIC_URL or "未暴露"
    }


@app.post("/upload")
async def api_upload(files: List[UploadFile] = File(...)):
    """上传文件并立即入库（API 直接调用，无预览确认步骤）"""
    results = []
    for file in files:
        try:
            # 保存临时文件
            temp_path = TEMP_DIR / f"{uuid.uuid4().hex}_{file.filename}"
            with open(temp_path, "wb") as f:
                content = await file.read()
                f.write(content)

            # 分类
            classification = classify_file(temp_path)

            # 复制到存储目录
            final_path = STORAGE_DIR / (str(uuid.uuid4())[:8] + "_" + file.filename)
            shutil.copy2(temp_path, final_path)

            # 入库
            doc_id = index_file(final_path, classification)

            # 清理临时文件
            temp_path.unlink(missing_ok=True)

            results.append({
                "status": "ok",
                "filename": file.filename,
                "doc_id": doc_id,
                "type": classification.get('type', '未知'),
                "tags": classification.get('tags', []),
                "summary": classification.get('summary', ''),
                "file_path": str(final_path),
                "download_url": f"/download/{doc_id}"
            })
        except Exception as e:
            results.append({
                "status": "error",
                "filename": file.filename,
                "error": str(e)
            })

    return JSONResponse(content={"results": results})


@app.post("/upload-preview")
async def api_upload_preview(files: List[UploadFile] = File(...)):
    """上传文件，分类但不入库（用于外部调用者预览确认）"""
    results = []
    for file in files:
        try:
            temp_path = TEMP_DIR / f"preview_{uuid.uuid4().hex}_{file.filename}"
            with open(temp_path, "wb") as f:
                content = await file.read()
                f.write(content)

            classification = classify_file(temp_path)
            sample = get_file_sample(temp_path)

            results.append({
                "temp_id": temp_path.name,
                "temp_path": str(temp_path),
                "filename": file.filename,
                "type": classification.get('type', '未知'),
                "tags": classification.get('tags', []),
                "summary": classification.get('summary', ''),
                "sample": sample[:200]
            })
        except Exception as e:
            results.append({
                "status": "error",
                "filename": file.filename,
                "error": str(e)
            })

    return JSONResponse(content={"results": results})


@app.post("/confirm-upload")
async def api_confirm_upload(data: str = Form(...)):
    """确认预览结果并入库（与 upload-preview 配对使用）

    data 格式: JSON 数组
    [{"temp_path":"...", "filename":"...", "type":"...", "tags":["..."], "summary":"..."}, ...]
    """
    try:
        items = json.loads(data)
        results = []
        for item in items:
            temp_path = Path(item["temp_path"])
            if not temp_path.exists():
                results.append({
                    "status": "error",
                    "filename": item["filename"],
                    "error": "临时文件不存在，可能已过期"
                })
                continue

            classification = {
                "type": item["type"],
                "tags": item["tags"],
                "summary": item["summary"]
            }
            final_path = STORAGE_DIR / (str(uuid.uuid4())[:8] + "_" + item["filename"])
            shutil.copy2(temp_path, final_path)
            doc_id = index_file(final_path, classification)
            temp_path.unlink(missing_ok=True)

            results.append({
                "status": "ok",
                "filename": item["filename"],
                "doc_id": doc_id,
                "file_path": str(final_path),
                "download_url": f"/download/{doc_id}"
            })

        return JSONResponse(content={"results": results})
    except Exception as e:
        raise HTTPException(400, detail=f"确认入库失败: {e}")


@app.get("/search")
async def search(q: str = Query(..., description="搜索关键词"), top_k: int = Query(5, description="返回结果数量")):
    """搜索知识库内容"""
    results = collection.query(query_texts=[q], n_results=top_k)
    if not results['documents'][0]:
        return {"answer": "没有相关内容。", "files": []}

    docs = results['documents'][0]
    metas = results['metadatas'][0]
    ids = results['ids'][0]

    # 构建回答上下文
    ctx_parts = []
    file_list = []
    for i, (doc, meta, fid) in enumerate(zip(docs, metas, ids)):
        ctx_parts.append(f"[文件{i+1}] {meta.get('filename', '未知')}\n摘要: {meta.get('summary', '')}\n内容: {doc[:300]}")
        dl_base = NGROK_PUBLIC_URL or _DOWNLOAD_BASE
        dl_url = f"{dl_base}/download/{fid}"
        file_list.append({
            "file_id": fid,
            "filename": meta.get('filename', '未知'),
            "type": meta.get('type', '未知'),
            "tags": meta.get('tags', ''),
            "summary": meta.get('summary', ''),
            "upload_time": meta.get('upload_time', ''),
            "file_path": meta.get('file_path', ''),
            "download_url": dl_url
        })

    ctx = "\n\n".join(ctx_parts)

    resp = openai_client.chat.completions.create(
        model="deepseek-chat",
        messages=[{"role": "user", "content": f"根据以下参考内容回答问题：\n\n{ctx}\n\n问题：{q}"}],
        temperature=0.3
    )

    return {
        "answer": resp.choices[0].message.content,
        "files": file_list
    }


@app.get("/download/{file_id}")
async def download_file(file_id: str):
    """通过文件 ID 下载文件"""
    result = collection.get(ids=[file_id], include=['metadatas'])
    if not result['ids']:
        raise HTTPException(404, "文件记录不存在")

    meta = result['metadatas'][0]
    file_path = meta.get('file_path', '')

    if not file_path or not Path(file_path).exists():
        raise HTTPException(404, "文件本体已不存在于磁盘上")

    # 读取文件并返回
    file_bytes = Path(file_path).read_bytes()
    filename = Path(file_path).name
    # 对中文文件名进行 RFC 5987 编码
    from urllib.parse import quote
    encoded_filename = quote(filename, safe='')
    return Response(
        content=file_bytes,
        media_type='application/octet-stream',
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}",
            "Content-Length": str(len(file_bytes))
        }
    )


@app.get("/files")
async def list_files():
    """列出所有已索引文件"""
    all_data = get_all_files()
    if not all_data['ids']:
        return {"files": []}

    dl_base = NGROK_PUBLIC_URL or _DOWNLOAD_BASE
    files = []
    for i, fid in enumerate(all_data['ids']):
        meta = all_data['metadatas'][i]
        files.append({
            "file_id": fid,
            "filename": meta.get('filename', '未知'),
            "type": meta.get('type', '未知'),
            "tags": meta.get('tags', ''),
            "summary": meta.get('summary', ''),
            "upload_time": meta.get('upload_time', ''),
            "file_size": meta.get('file_size', 0),
            "file_path": meta.get('file_path', ''),
            "download_url": f"{dl_base}/download/{fid}"
        })

    files.sort(key=lambda x: x.get('upload_time', ''), reverse=True)
    return {"files": files}


@app.put("/files/{file_id}")
async def update_file_meta(file_id: str,
                           file_type: str = Form(...),
                           tags: str = Form(...),
                           summary: str = Form(...)):
    """更新文件元数据（分类、标签、摘要）"""
    ok = update_metadata(file_id, file_type, tags, summary)
    if not ok:
        raise HTTPException(500, "更新失败")
    return {"status": "ok", "file_id": file_id}


@app.delete("/files/{file_id}")
async def delete_file_record(file_id: str):
    """删除文件索引记录"""
    ok = delete_document(file_id)
    if not ok:
        raise HTTPException(500, "删除失败")
    return {"status": "ok", "file_id": file_id}


@app.post("/text-to-doc")
async def api_text_to_doc(text: str = Form(...), label: str = Form("")):
    """文字转文档并入库"""
    result = text_to_doc(text, label)
    if not result["success"]:
        raise HTTPException(500, result.get("error", "处理失败"))
    return {
        "status": "ok",
        "doc_id": result["doc_id"],
        "filename": result["filename"],
        "title": result["title"],
        "summary": result["summary"],
        "download_url": f"/download/{result['doc_id']}"
    }


# ==================== Gradio UI ====================

def gradio_upload_files(files):
    """Gradio 多文件上传：分类、展示可编辑表格（不入库）"""
    if not files:
        return [], [], "请拖拽文件"

    # files 可以是列表（多文件）或单个文件
    if not isinstance(files, list):
        files = [files]

    state_data = []
    table_data = []

    for i, f in enumerate(files):
        try:
            # Gradio 的 File 组件返回的对象有 .name 属性（临时文件路径）
            src_path = Path(f.name if hasattr(f, 'name') else f)
            filename = src_path.name

            # 复制到 temp_uploads 防止 Gradio 清理
            temp_path = TEMP_DIR / f"gradio_{uuid.uuid4().hex}_{filename}"
            shutil.copy2(src_path, temp_path)

            # 分类
            classification = classify_file(temp_path)

            state_data.append({
                "temp_path": str(temp_path),
                "filename": filename
            })
            table_data.append([
                filename,
                classification.get('type', '未知'),
                ", ".join(classification.get('tags', [])),
                classification.get('summary', '')
            ])

            # 间隔 0.5s 避免 API 频率限制
            if i < len(files) - 1:
                time.sleep(0.5)

        except Exception as e:
            table_data.append([
                f.name if hasattr(f, 'name') else str(f),
                "分类失败",
                "",
                f"错误: {e}"
            ])

    return state_data, table_data, f"已上传 {len(files)} 个文件，请核对分类结果后点击确认入库"


def gradio_confirm_upload(state_data, table_data):
    """确认分类结果并入库"""
    # 新版 Gradio Dataframe 返回 pandas DataFrame，需转 list
    try:
        import pandas as pd
        if isinstance(table_data, pd.DataFrame):
            table_data = table_data.values.tolist()
    except ImportError:
        pass

    if not table_data or not state_data:
        return "⚠️ 没有待入库的文件"

    success_count = 0
    fail_count = 0
    messages = []

    for i, row in enumerate(table_data):
        try:
            filename, file_type, tags_str, summary = row[:4]
            tags = [t.strip() for t in tags_str.split(",") if t.strip()]

            if i >= len(state_data):
                messages.append(f"❌ {filename}: 状态数据丢失")
                fail_count += 1
                continue

            temp_path = Path(state_data[i]["temp_path"])

            if not temp_path.exists():
                messages.append(f"❌ {filename}: 临时文件已过期，请重新上传")
                fail_count += 1
                continue

            classification = {
                "type": file_type,
                "tags": tags,
                "summary": summary
            }

            # 复制到存储目录
            final_path = STORAGE_DIR / (str(uuid.uuid4())[:8] + "_" + filename)
            shutil.copy2(temp_path, final_path)

            # 入库
            doc_id = index_file(final_path, classification)

            # 删除临时文件
            temp_path.unlink(missing_ok=True)

            messages.append(f"✅ {filename} → {file_type}")
            success_count += 1

        except Exception as e:
            messages.append(f"❌ {row[0] if row else '未知'}: {e}")
            fail_count += 1

    result = f"✅ 入库完成: {success_count} 成功"
    if fail_count:
        result += f", {fail_count} 失败"
    result += "\n" + "\n".join(messages[:10])
    if len(messages) > 10:
        result += f"\n... 共 {len(messages)} 条记录"

    return result


def gradio_chat(message, history):
    """Gradio 对话查询"""
    if not message:
        return "", history

    results = collection.query(query_texts=[message], n_results=5)
    if not results['documents'][0]:
        reply = "知识库中没有相关内容。"
    else:
        docs = results['documents'][0]
        metas = results['metadatas'][0]
        ids = results['ids'][0]

        dl_base = NGROK_PUBLIC_URL or _DOWNLOAD_BASE

        ctx_parts = []
        file_links = []
        for i, (doc, meta, fid) in enumerate(zip(docs, metas, ids)):
            upload_time = meta.get('upload_time', '')[:19] if meta.get('upload_time') else '未知'
            dl_url = f"{dl_base}/download/{fid}"
            ctx_parts.append(
                f"[文件{i+1}] 名称: {meta.get('filename', '未知')}\n"
                f"类型: {meta.get('type', '未知')} | 标签: {meta.get('tags', '')}\n"
                f"上传时间: {upload_time}\n"
                f"路径: {meta.get('file_path', '')}\n"
                f"内容:\n{doc}"
            )
            file_links.append(
                f"{meta.get('filename', '未知')} "
                f"({meta.get('type', '未知')}) "
                f"下载: {dl_url}"
            )

        ctx = "\n\n---\n\n".join(ctx_parts)
        files_str = "\n".join(file_links)

        prompt = f"""你是一个知识库助手。请根据下面的参考文件内容回答问题。

【重要规则】
- 优先基于文件内容给出具体、准确的答案
- 如果用户问文件在哪、怎么下载，直接给出完整路径和下载链接（如下方"文件来源"所示）
- 如果内容不足以回答，诚实说明，不要编造

【参考文件内容】
{ctx}

【文件来源与下载链接】
{files_str}

【用户问题】
{message}

【回答】"""

        resp = openai_client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1
        )
        reply = resp.choices[0].message.content
        reply += f"\n\n---\n📥 下载链接: " + "  |  ".join(file_links)

    if history is None:
        history = []
    history.append({"role": "user", "content": message})
    history.append({"role": "assistant", "content": reply})
    return "", history


def gradio_text_to_doc(text, label):
    """Gradio 文本转文档"""
    if not text.strip():
        return "⚠️ 请输入文字内容"

    result = text_to_doc(text, label)
    if not result["success"]:
        return f"❌ 处理失败: {result.get('error', '未知错误')}"

    return (
        f"✅ 文档已生成并入库\n\n"
        f"📄 文件名: {result['filename']}\n"
        f"📌 标题: {result['title']}\n"
        f"📝 摘要: {result['summary']}\n"
        f"🔗 下载: /download/{result['doc_id']}\n"
        f"💡 现在可以在对话中查询到这篇文档的内容"
    )


def gradio_refresh_files():
    """刷新文件管理列表"""
    all_data = get_all_files()
    if not all_data['ids']:
        return [], "📭 暂无已索引文件"

    table = []
    for i, fid in enumerate(all_data['ids']):
        meta = all_data['metadatas'][i]
        upload_time = meta.get('upload_time', '')
        if upload_time:
            upload_time = upload_time[:19].replace('T', ' ')
        table.append([
            fid[:8] + "...",
            meta.get('filename', '未知'),
            meta.get('type', '未知'),
            meta.get('tags', ''),
            meta.get('summary', '')[:40],
            upload_time
        ])

    table.sort(key=lambda x: x[5], reverse=True)
    return table, f"共 {len(table)} 个文件"


def gradio_select_file(evt: gr.SelectData, files_table):
    """从文件管理表格选择一行，填充到编辑表单"""
    try:
        import pandas as pd
        if isinstance(files_table, pd.DataFrame):
            files_table = files_table.values.tolist()
    except ImportError:
        pass

    if not files_table or evt.index[0] >= len(files_table):
        return "", "", "", "", ""

    row = files_table[evt.index[0]]

    # 根据文件名查找完整 ID
    all_data = get_all_files()
    target_fid = ""
    target_meta = None
    for i, fid in enumerate(all_data['ids']):
        if all_data['metadatas'][i].get('filename', '') == row[1]:
            target_fid = fid
            target_meta = all_data['metadatas'][i]
            break

    if not target_fid:
        return "", "", "", "", ""

    file_type = target_meta.get('type', '文本笔记')
    tags = target_meta.get('tags', '')
    summary = target_meta.get('summary', '')

    return target_fid, file_type, tags, summary, f"已选中: {row[1]}"


def gradio_save_edit(file_id, file_type, tags, summary):
    """保存文件管理中的编辑"""
    if not file_id:
        return "⚠️ 请先选择一个文件"
    ok = update_metadata(file_id, file_type, tags, summary)
    if ok:
        return f"✅ 已更新文件 {file_id[:8]}..."
    return "❌ 更新失败"


def gradio_delete_file(file_id):
    """删除文件记录"""
    if not file_id:
        return "⚠️ 请先选择一个文件"
    ok = delete_document(file_id)
    if ok:
        return f"✅ 已删除文件记录 {file_id[:8]}..."
    return "❌ 删除失败"


def build_gradio_ui():
    """构建 Gradio UI"""
    with gr.Blocks(title="私人AI文件管家") as demo:
        # ========== 顶部标题 ==========
        with gr.Row():
            gr.Markdown(
                """
                # 👤 私人AI文件管家
                **本地知识库 · 智能分类 · 语义检索 · 公网访问**
                """
            )

        # ========== 状态栏 ==========
        status_text = gr.Markdown(
            f"> 🟢 API 运行中 | 🟢 模型就绪"
            + (f" | 🌐 公网地址: `{NGROK_PUBLIC_URL}`" if NGROK_PUBLIC_URL else " | 🔒 本地模式")
        )

        # ========== 主布局：左侧 Tab 面板 + 右侧 Chat ==========
        with gr.Row():
            # ---- 左侧 Tab 面板 ----
            with gr.Column(scale=2):
                with gr.Tabs():
                    # ===== Tab 1: 文件上传 =====
                    with gr.TabItem("📤 文件上传", id="tab_upload"):
                        gr.Markdown("### 拖拽或选择文件，自动分类后确认入库")
                        with gr.Row():
                            file_input = gr.File(
                                label="拖拽文件到此处（支持多选）",
                                file_count="multiple",
                                file_types=[".txt", ".md", ".py", ".json", ".yaml", ".yml",
                                            ".ini", ".cfg", ".log", ".docx", ".pdf",
                                            ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".zip", ".rar"]
                            )
                        upload_btn = gr.Button("🔄 自动分类", variant="primary")
                        upload_status = gr.Textbox(label="操作状态", interactive=False)

                        gr.Markdown("#### ✏️ 核对并编辑分类结果（双击单元格可修改）")
                        pending_table = gr.Dataframe(
                            headers=["文件名", "类型", "标签(逗号分隔)", "摘要"],
                            datatype=["str", "str", "str", "str"],
                            value=[],
                            interactive=True,
                            label="待入库文件列表"
                        )
                        confirm_btn = gr.Button("✅ 确认全部入库", variant="primary", visible=True)

                        # 内部状态（隐藏，存临时路径）
                        upload_state = gr.State([])

                        # 事件绑定
                        upload_btn.click(
                            fn=gradio_upload_files,
                            inputs=file_input,
                            outputs=[upload_state, pending_table, upload_status]
                        )
                        confirm_btn.click(
                            fn=gradio_confirm_upload,
                            inputs=[upload_state, pending_table],
                            outputs=upload_status
                        ).then(
                            fn=lambda: ([], []),
                            inputs=None,
                            outputs=[upload_state, pending_table]
                        )

                    # ===== Tab 2: 文本转文档 =====
                    with gr.TabItem("📝 文本转文档", id="tab_text2doc"):
                        gr.Markdown("### 将一段文字转为文档，自动总结并存入库")
                        with gr.Column():
                            text_input = gr.Textbox(
                                label="输入文字内容",
                                placeholder="粘贴或输入文字，AI 会自动总结成文档...",
                                lines=10
                            )
                            with gr.Row():
                                label_input = gr.Textbox(
                                    label="标记（可选）",
                                    placeholder="例如: API密钥说明、项目总结、会议记录...",
                                    scale=2
                                )
                                gen_btn = gr.Button("🔄 生成文档并入库", variant="primary", scale=1)
                            gen_output = gr.Textbox(label="操作结果", lines=6, interactive=False)

                        gen_btn.click(
                            fn=gradio_text_to_doc,
                            inputs=[text_input, label_input],
                            outputs=gen_output
                        )

                    # ===== Tab 3: 文件管理 =====
                    with gr.TabItem("📂 文件管理", id="tab_files"):
                        gr.Markdown("### 浏览、编辑、删除已索引文件")
                        refresh_btn = gr.Button("🔄 刷新列表", variant="secondary")
                        files_table = gr.Dataframe(
                            headers=["ID", "文件名", "类型", "标签", "摘要", "上传时间"],
                            datatype=["str", "str", "str", "str", "str", "str"],
                            value=[],
                            interactive=False,
                            label="全部已索引文件（点击行选中编辑）",
                            row_count=20
                        )
                        file_count = gr.Markdown("📭 暂无已索引文件")

                        gr.Markdown("---")
                        gr.Markdown("#### ✏️ 编辑选中文件")
                        with gr.Row():
                            edit_fid = gr.Textbox(label="文件ID（自动填充）", interactive=False, scale=1)
                            edit_type = gr.Dropdown(
                                label="文件类型",
                                choices=FILE_TYPES,
                                scale=1
                            )
                            edit_tags = gr.Dropdown(
                                label="标签（多选）",
                                choices=FILE_TAGS,
                                multiselect=True,
                                scale=2
                            )
                        edit_summary = gr.Textbox(label="摘要", lines=2)
                        with gr.Row():
                            edit_status = gr.Textbox(label="操作结果", interactive=False, scale=3)
                            save_btn = gr.Button("💾 保存修改", variant="primary", scale=1)
                            delete_btn = gr.Button("🗑️ 删除记录", variant="stop", scale=1)

                        # 事件绑定
                        refresh_btn.click(
                            fn=gradio_refresh_files,
                            outputs=[files_table, file_count]
                        )
                        files_table.select(
                            fn=gradio_select_file,
                            inputs=files_table,
                            outputs=[edit_fid, edit_type, edit_tags, edit_summary, edit_status]
                        )
                        save_btn.click(
                            fn=gradio_save_edit,
                            inputs=[edit_fid, edit_type, edit_tags, edit_summary],
                            outputs=edit_status
                        ).then(
                            fn=gradio_refresh_files,
                            outputs=[files_table, file_count]
                        )
                        delete_btn.click(
                            fn=gradio_delete_file,
                            inputs=edit_fid,
                            outputs=edit_status
                        ).then(
                            fn=gradio_refresh_files,
                            outputs=[files_table, file_count]
                        )

            # ---- 右侧 Chat 面板 ----
            with gr.Column(scale=2):
                gr.Markdown("### 💬 对话查询")
                chatbot = gr.Chatbot(
                    label="知识库智能问答",
                    height=500,
                    # bubble_limit removed in Gradio 6.0
                )
                with gr.Row():
                    msg = gr.Textbox(
                        label="输入问题",
                        placeholder="用自然语言提问，例如：帮我找一下昨天上传的密钥文件...",
                        scale=5
                    )
                    send_btn = gr.Button("发送", variant="primary", scale=1)
                clear_btn = gr.Button("清空对话", variant="secondary", size="sm")

                # 事件绑定
                send_btn.click(fn=gradio_chat, inputs=[msg, chatbot], outputs=[msg, chatbot])
                msg.submit(fn=gradio_chat, inputs=[msg, chatbot], outputs=[msg, chatbot])
                clear_btn.click(fn=lambda: ([], ""), outputs=[chatbot, msg])

    return demo


# ==================== 启动入口 ====================

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="MyKB 私人知识库")
    parser.add_argument("--expose", action="store_true", help="通过 ngrok 暴露 API 到公网")
    parser.add_argument("--port", type=int, default=8000, help="FastAPI 端口 (默认 8000)")
    parser.add_argument("--gradio-port", type=int, default=7860, help="Gradio 端口 (默认 7860)")
    args = parser.parse_args()

    # 启动 ngrok（如果指定）
    if args.expose:
        start_ngrok(args.port)

    # 启动 FastAPI（后台线程）
    print(f"\n[启动] FastAPI (端口 {args.port})...")
    t = threading.Thread(
        target=lambda: uvicorn.run(app, host="0.0.0.0", port=args.port, log_level="info"),
        daemon=True
    )
    t.start()

    # 设置下载基础地址
    _DOWNLOAD_BASE = f"http://127.0.0.1:{args.port}"

    # 启动 Gradio
    print(f"[启动] Gradio UI (端口 {args.gradio_port})...")
    demo = build_gradio_ui()
    demo.launch(
        server_name="127.0.0.1",
        server_port=args.gradio_port,
        share=False,
        show_error=True
    )
